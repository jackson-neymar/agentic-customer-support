# customer_support_chat/app/services/rag/document_loader.py

"""
文档加载器模块 - 支持 PDF, DOCX, Markdown, TXT
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass
import hashlib
import re

# =====================================
# 1. 基础类定义
# =====================================

@dataclass
class LoadedDocument:
    """加载的文档（区分于 RetrievalResult）"""
    content: str
    metadata: Dict[str, Any]
    
    @property
    def id(self) -> str:
        """生成文档唯一 ID"""
        return hashlib.md5(self.content.encode()).hexdigest()[:16]


class BaseDocumentLoader(ABC):
    """文档加载器基类"""
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
    
    @abstractmethod
    def load(self) -> List[LoadedDocument]:
        """加载文档"""
        pass
    
    def _get_base_metadata(self) -> Dict[str, Any]:
        """获取文件基础元数据"""
        stat = self.file_path.stat()
        return {
            'source': str(self.file_path),
            'filename': self.file_path.name,
            'file_type': self.file_path.suffix,
            'file_size': stat.st_size,
        }


# =====================================
# 2. 具体加载器实现
# =====================================

class PDFLoader(BaseDocumentLoader):
    """PDF 文档加载器"""
    
    def load(self) -> List[LoadedDocument]:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("请安装: pip install pypdf")
        
        documents = []
        reader = PdfReader(str(self.file_path))
        
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text().strip()
            if not text:
                continue
            
            metadata = self._get_base_metadata()
            metadata.update({
                'page': page_num,
                'total_pages': len(reader.pages),
            })
            
            documents.append(LoadedDocument(content=text, metadata=metadata))
        
        return documents


class DocxLoader(BaseDocumentLoader):
    """Word 文档加载器"""
    
    def load(self) -> List[LoadedDocument]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装: pip install python-docx")
        
        doc = DocxDocument(str(self.file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n\n'.join(paragraphs)
        
        metadata = self._get_base_metadata()
        metadata['paragraph_count'] = len(paragraphs)
        
        return [LoadedDocument(content=content, metadata=metadata)]


class MarkdownLoader(BaseDocumentLoader):
    """Markdown 文档加载器"""
    
    def load(self) -> List[LoadedDocument]:
        with open(self.file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取标题
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        
        metadata = self._get_base_metadata()
        metadata['title'] = title_match.group(1) if title_match else ''
        
        return [LoadedDocument(content=content, metadata=metadata)]


class TextLoader(BaseDocumentLoader):
    """纯文本加载器"""
    
    def load(self) -> List[LoadedDocument]:
        with open(self.file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        metadata = self._get_base_metadata()
        return [LoadedDocument(content=content, metadata=metadata)]


# =====================================
# 3. 文档加载器工厂
# =====================================

class DocumentLoaderFactory:
    """文档加载器工厂"""
    
    LOADERS = {
        '.pdf': PDFLoader,
        '.docx': DocxLoader,
        '.doc': DocxLoader,
        '.md': MarkdownLoader,
        '.markdown': MarkdownLoader,
        '.txt': TextLoader,
    }
    
    @classmethod
    def load_document(cls, file_path: str) -> List[LoadedDocument]:
        """加载单个文档"""
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix not in cls.LOADERS:
            raise ValueError(
                f"不支持的文件类型: {suffix}\n"
                f"支持: {', '.join(cls.LOADERS.keys())}"
            )
        
        loader = cls.LOADERS[suffix](file_path)
        return loader.load()
    
    @classmethod
    def load_directory(
        cls, 
        dir_path: str, 
        recursive: bool = True,
        verbose: bool = True
    ) -> List[LoadedDocument]:
        """加载目录中的所有文档"""
        path = Path(dir_path)
        if not path.is_dir():
            raise ValueError(f"不是有效目录: {dir_path}")
        
        documents = []
        pattern = '**/*' if recursive else '*'
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in cls.LOADERS:
                try:
                    docs = cls.load_document(str(file_path))
                    documents.extend(docs)
                    if verbose:
                        print(f"✅ 已加载: {file_path.name} ({len(docs)} 文档)")
                except Exception as e:
                    if verbose:
                        print(f"⚠️ 加载失败 {file_path.name}: {e}")
        
        return documents


# =====================================
# 4. 文本分块器
# =====================================

class SimpleTextSplitter:
    """简单的文本分块器"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_documents(self, documents: List[LoadedDocument]) -> List[LoadedDocument]:
        """分块文档"""
        chunked_docs = []
        
        for doc in documents:
            chunks = self._split_text(doc.content)
            
            for i, chunk in enumerate(chunks):
                metadata = doc.metadata.copy()
                metadata.update({
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                })
                
                chunked_docs.append(LoadedDocument(
                    content=chunk,
                    metadata=metadata
                ))
        
        return chunked_docs
    
    def _split_text(self, text: str) -> List[str]:
        """简单分块：按段落和句子"""
        # 先按双换行符分段
        paragraphs = text.split('\n\n')
        
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) <= self.chunk_size:
                current_chunk += para + '\n\n'
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + '\n\n'
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks